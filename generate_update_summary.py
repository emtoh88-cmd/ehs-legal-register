#!/usr/bin/env python3
"""
generate_update_summary.py — turn the changes sso_refresh.py has found into a
personal "Update Summary" report, in the shape of the quarterly update
summaries a paid EHS compliance subscription would send.

Source of truth is register.json's meta.pending_review, written by
sso_refresh.py whenever it finds an instrument whose SSO version date or
amendment list has moved since the register last recorded it. Nothing here
is fetched from SSO itself -- this only formats what refresh already found.

Usage:
  python generate_update_summary.py                 # write reports/update-summary-YYYY-MM.docx
  python generate_update_summary.py --apply          # also fold the changes into register.json's
                                                       # history log and clear pending_review
  python generate_update_summary.py --out FILE.docx  # choose the output path
"""

import argparse, json, os
from datetime import date, datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REGISTER = "register.json"
REPORTS_DIR = "reports"

MONTHS = ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
          "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]


def strip_num(cat):
    """'5 Safety and Health' -> 'Safety and Health', to match the history log's style."""
    parts = cat.split(" ", 1)
    return parts[1] if parts and parts[0].replace(".", "").isdigit() else cat


def build_category_map(reg):
    cat = {}
    for c in reg["contents"]:
        if c.get("citation"):
            cat[c["citation"]] = strip_num(c["cat"])
    return cat


def parse_version_date(s):
    """'1 Feb 2024' -> date(2024, 2, 1); returns None if unparseable/blank."""
    if not s:
        return None
    try:
        return datetime.strptime(s.strip(), "%d %b %Y").date()
    except ValueError:
        return None


def quarter_of(d):
    return f"{d.year}-{(d.month - 1) // 3 + 1}Q"


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def build_docx(reg, changes, out_path):
    cat_map = build_category_map(reg)
    for c in changes:
        c["category"] = cat_map.get(c["citation"], "Uncategorised")
        c["version_date"] = parse_version_date(c.get("to") or c.get("from"))

    # group by month, oldest first -- undated changes go in a trailing group
    dated = sorted((c for c in changes if c["version_date"]), key=lambda c: c["version_date"])
    undated = [c for c in changes if not c["version_date"]]
    groups = {}
    for c in dated:
        groups.setdefault((c["version_date"].year, c["version_date"].month), []).append(c)
    if undated:
        groups[(9999, 0)] = undated

    today = date.today()
    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Cm(1.8)

    title = doc.add_heading("EHS Legislation Update Summary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f"Personal register · generated {today.isoformat()} · "
        f"{len(changes)} instrument{'s' if len(changes) != 1 else ''} flagged by sso_refresh.py"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(10)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Cm(1.3), Cm(3.2), Cm(9.0), Cm(3.0)]
    headers = ["S/No", "Category", "Legislation updated", "Action"]
    for i, (cell, w, h) in enumerate(zip(table.rows[0].cells, widths, headers)):
        cell.width = w
        set_cell_text(cell, h, bold=True, size=10)
        shade_cell(cell, "D9D9D9")

    n = 0
    for (y, m) in sorted(groups.keys()):
        row = table.add_row().cells
        label = "Undated / no version date recorded" if m == 0 else f"{MONTHS[m - 1]} {y}"
        row[0].merge(row[1]).merge(row[2]).merge(row[3])
        set_cell_text(row[0], label, bold=True, size=10)
        shade_cell(row[0], "F2F2F2")

        for c in groups[(y, m)]:
            n += 1
            row = table.add_row().cells
            for cell, w in zip(row, widths):
                cell.width = w
            set_cell_text(row[0], str(n))
            set_cell_text(row[1], c["category"])

            lines = [c["citation"]]
            if c.get("from") and c.get("to") and c["from"] != c["to"]:
                lines.append(f"Version date: {c['from']} → {c['to']}")
            elif c.get("to"):
                lines.append(f"Version date: {c['to']}")
            if c.get("new"):
                lines.append("Amended by: " + ", ".join(c["new"]))
            cell3 = row[2]
            cell3.text = ""
            p3 = cell3.paragraphs[0]
            p3.add_run(lines[0]).bold = True
            for extra in lines[1:]:
                p = cell3.add_paragraph()
                p.add_run(extra).font.size = Pt(9)

            set_cell_text(row[3], "")  # left for the reviewer to fill in

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "This report lists version-date and amendment changes detected by sso_refresh.py against "
        "Singapore Statutes Online. It is not a substitute for reading the amending instrument -- "
        "follow the SSO link in the register to confirm applicability before closing out the "
        "\"Action\" column."
    ).italic = True
    note.runs[0].font.size = Pt(8)

    copyright_p = doc.add_paragraph()
    copyright_p.add_run(reg["meta"].get("notice", "")).font.size = Pt(8)

    doc.save(out_path)
    return n


def apply_to_history(reg, changes):
    """Fold reviewed changes into meta.history the way the vendor's Update History sheet works."""
    history = reg.setdefault("history", [])
    next_no = max((int(h["no"]) for h in history if str(h.get("no", "")).isdigit()), default=0)
    cat_map = build_category_map(reg)

    for c in changes:
        d = parse_version_date(c.get("to") or c.get("from"))
        next_no += 1
        history.append({
            "quarter": quarter_of(d) if d else "",
            "no": str(next_no),
            "category": cat_map.get(c["citation"], "Uncategorised"),
            "title": c["citation"],
            "wef": d.isoformat() if d else "",
            "vide": ", ".join(c.get("new", [])) or (c.get("to") or ""),
        })

    reg["meta"]["pending_review"] = []
    reg["meta"]["counts"]["history"] = len(history)
    if changes:
        latest = max((parse_version_date(c.get("to") or c.get("from")) for c in changes
                      if parse_version_date(c.get("to") or c.get("from"))), default=None)
        if latest:
            reg["meta"]["source_snapshot"] = f"{latest.year}-{latest.month:02d}"


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--out", help="output .docx path (default: reports/update-summary-YYYY-MM.docx)")
    p.add_argument("--apply", action="store_true",
                    help="also append the changes to register.json's history and clear pending_review")
    args = p.parse_args()

    if not os.path.exists(REGISTER):
        print(f"{REGISTER} not found"); return 1
    reg = json.load(open(REGISTER, encoding="utf-8"))

    changes = reg.get("meta", {}).get("pending_review", [])
    if not changes:
        print("No pending changes in register.json (meta.pending_review is empty). "
              "Run sso_refresh.py first.")
        return 0

    os.makedirs(REPORTS_DIR, exist_ok=True)
    out_path = args.out or os.path.join(REPORTS_DIR, f"update-summary-{date.today():%Y-%m}.docx")

    n = build_docx(reg, changes, out_path)
    print(f"wrote {out_path} ({n} instruments)")

    if args.apply:
        apply_to_history(reg, changes)
        tmp = REGISTER + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(reg, f, separators=(",", ":"), ensure_ascii=False)
        os.replace(tmp, REGISTER)
        print(f"applied {len(changes)} changes to {REGISTER}'s history log")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
